import io
import json
import pandas as pd

def extract_text_from_pdf(file_bytes):
    """Extracts all text from a PDF file using pypdf."""
    import pypdf  # Pure Python PDF reader to prevent segmentation faults
    text = ""
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text_from_docx(file_bytes):
    """Extracts all text from a Word (.docx) file using python-docx."""
    import docx  # python-docx
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs_text = [p.text for p in doc.paragraphs]
    
    # Also extract text from tables
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text for cell in row.cells]
            table_text.append(" | ".join(row_cells))
            
    return "\n".join(paragraphs_text + table_text)

def extract_text_from_excel_csv(file_bytes, filename):
    """Reads Excel/CSV files and converts the top contents to a text representation."""
    if filename.endswith('.csv'):
        df = pd.read_csv(io.BytesIO(file_bytes))
    else:
        df = pd.read_excel(io.BytesIO(file_bytes))
        
    # Get shape and basic details
    summary = f"Archivo Excel/CSV: {filename}\n"
    summary += f"Dimensiones: {df.shape[0]} filas, {df.shape[1]} columnas\n"
    summary += f"Columnas: {', '.join(df.columns.tolist())}\n"
    summary += "Primeras 50 filas de datos:\n"
    # Convert first 50 rows to markdown or CSV table format
    summary += df.head(50).to_csv(index=False)
    return summary

def parse_business_file(file_bytes, filename):
    """Dispatches parsing depending on file extension."""
    ext = filename.split('.')[-1].lower()
    
    if ext == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif ext in ['doc', 'docx']:
        return extract_text_from_docx(file_bytes)
    elif ext in ['csv', 'xlsx', 'xls']:
        return extract_text_from_excel_csv(file_bytes, filename)
    else:
        # Fallback to plain text reading
        try:
            return file_bytes.decode('utf-8')
        except Exception:
            raise ValueError(f"Extensión de archivo .{ext} no soportada.")
def clean_to_ascii(text):
    """
    Replaces Spanish accented characters with their ASCII equivalents and strips other non-ASCII chars.
    This prevents UnicodeEncodeErrors (e.g. 'ascii' codec can't encode...) on systems with restricted locales or in gRPC/protobuf serialization.
    """
    import unicodedata
    if not text:
        return ""
    # Normalize unicode to separate characters from their accents (NFD)
    nfd_form = unicodedata.normalize('NFD', text)
    only_ascii = "".join([c for c in nfd_form if not unicodedata.combining(c)])
    
    # Manual replacement for Spanish 'ñ'/'Ñ' and other common letters that NFD doesn't split to base ascii characters
    replacements = {
        'ñ': 'n', 'Ñ': 'N',
        'á': 'a', 'é': 'e', 'í': 'i', 'ó': 'o', 'ú': 'u',
        'Á': 'A', 'É': 'E', 'Í': 'I', 'Ó': 'O', 'Ú': 'U',
        'ü': 'u', 'Ü': 'U'
    }
    for orig, rep in replacements.items():
        only_ascii = only_ascii.replace(orig, rep)
        
    return only_ascii.encode('ascii', errors='ignore').decode('ascii')

def extract_parameters_via_openai(document_text, filename, api_key, model_name="gpt-4o-mini"):
    """
    Calls the OpenAI API to extract financial parameters from raw text.
    Ensures response conforms to a clean, flat JSON schema matching the simulation requirements.
    """
    if isinstance(api_key, str):
        api_key = api_key.strip()

    import openai
    import json
    
    client = openai.OpenAI(api_key=api_key)
    
    prompt = f"""
    Eres un analista financiero experto en gestion de riesgos y simulacion de Monte Carlo. 
    Se te proporciona el contenido extraido de un archivo de negocio real ({filename}).
    Tu tarea es analizar la informacion y extraer los parametros base y de incertidumbre para una simulacion financiera de 4 meses.
    
    El modelo de simulacion de 4 meses asume:
    - 3 meses de temporada baja (sujetos a contraccion de demanda)
    - 1 mes de temporada alta (sujeto a un pico o repunte de demanda)
    - Riesgo de aumento en el costo de un insumo o materia prima critica (que representa un % del costo de ventas o COGS).
    
    Por favor, lee el documento y extrae u optimiza los siguientes parametros. Si el documento no especifica alguno de los valores financieros, utiliza tu conocimiento sectorial y el contexto de la empresa para ESTIMAR montos y desviaciones logicas (especifica la justificacion de estas estimaciones en la descripcion).
    
    Debes devolver un objeto JSON plano que contenga EXACTAMENTE los siguientes campos:
    
    1. "scenario_title": Un nombre corto y descriptivo del negocio/proyecto (max 50 caracteres).
    2. "scenario_description": Un resumen ejecutivo (max 300 caracteres) que describa de que trata la empresa, que insumo critico es el volatil y el contexto del mercado geografico o sectorial.
    3. "baseline_revenue": Ingresos mensuales promedio base del negocio en condiciones normales. Debe ser un numero float.
    4. "baseline_cogs": Costo mensual de ventas base (COGS) promedio del negocio. Debe ser un numero float.
    5. "raw_material_share": Proporcion de la materia prima/insumo critico dentro del costo de ventas total (ej. 0.35 si representa el 35% del COGS). Debe ser un float entre 0.05 y 0.90.
    6. "fixed_costs": Gastos fijos mensuales promedio (renta, nomina fija, servicios, seguros). Debe ser un numero float.
    7. "raw_material_risk_prob": Probabilidad de que el precio del insumo critico aumente bruscamente debido a factores de riesgo (clima, logistica, aranceles) (ej. 0.30 para un 30% de probabilidad). Debe ser un float entre 0.0 y 1.0.
    8. "raw_material_risk_increase": Porcentaje estimado de aumento en el costo del insumo en caso de crisis (ej. 0.50 si sube un 50% de precio). Debe ser un float entre 0.0 y 2.0.
    9. "low_season_contraction_mean": Contraccion promedio esperada de la demanda general en los 3 meses de temporada baja (ej. 0.20 para una reduccion del 20% en ventas). Debe ser un float entre 0.0 y 1.0.
    10. "low_season_contraction_std": Volatilidad o desviacion estandar de la contraccion de la demanda en temporada baja. Valor sugerido entre 0.02 y 0.15 (por defecto 0.05).
    11. "high_season_spike_mean": Repunte o incremento promedio esperado de la demanda en el mes de temporada alta (ej. 0.60 para un incremento de 60%). Debe ser un float entre 0.0 y 2.0.
    12. "high_season_spike_std": Volatilidad o desviacion estandar del incremento en temporada alta. Valor sugerido entre 0.05 y 0.25 (por defecto 0.10).
    
    Contenido del documento:
    ---
    {document_text}
    ---
    
    Devuelve UNICAMENTE el objeto JSON plano solicitado.
    """
    
    # Sanitize the prompt string to prevent ascii encoding issues in API clients
    prompt = clean_to_ascii(prompt)
    
    response = client.chat.completions.create(
        model=model_name,
        messages=[
            {"role": "system", "content": "Eres un analista financiero experto en analisis de riesgo."},
            {"role": "user", "content": prompt}
        ],
        response_format={"type": "json_object"}
    )
    
    try:
        data = json.loads(response.choices[0].message.content)
        return data
    except Exception as e:
        cleaned_text = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
        return json.loads(cleaned_text)

def extract_parameters_via_gemini(document_text, filename, api_key, model_name="gemini-1.5-flash"):
    if isinstance(api_key, str):
        api_key = api_key.strip()

    # Route to OpenAI if an OpenAI model is specified
    if model_name.startswith("gpt-") or model_name.startswith("o1-") or model_name.startswith("o3-"):
        return extract_parameters_via_openai(document_text, filename, api_key, model_name)

    """
    Calls the Gemini API to extract financial parameters from raw text.
    Ensures response conforms to a clean, flat JSON schema matching the simulation requirements.
    """
    import google.generativeai as genai
    # Configure Gemini API
    genai.configure(api_key=api_key)
    
    prompt = f"""
    Eres un analista financiero experto en gestion de riesgos y simulacion de Monte Carlo. 
    Se te proporciona el contenido extraido de un archivo de negocio real ({filename}).
    Tu tarea es analizar la informacion y extraer los parametros base y de incertidumbre para una simulacion financiera de 4 meses.
    
    El modelo de simulacion de 4 meses asume:
    - 3 meses de temporada baja (sujetos a contraccion de demanda)
    - 1 mes de temporada alta (sujeto a un pico o repunte de demanda)
    - Riesgo de aumento en el costo de un insumo o materia prima critica (que representa un % del costo de ventas o COGS).
    
    Por favor, lee el documento y extrae u optimiza los siguientes parametros. Si el documento no especifica alguno de los valores financieros, utiliza tu conocimiento sectorial y el contexto de la empresa para ESTIMAR montos y desviaciones logicas (especifica la justificacion de estas estimaciones en la descripcion).
    
    Debes devolver un objeto JSON plano que contenga EXACTAMENTE los siguientes campos:
    
    1. "scenario_title": Un nombre corto y descriptivo del negocio/proyecto (max 50 caracteres).
    2. "scenario_description": Un resumen ejecutivo (max 300 caracteres) que describa de que trata la empresa, que insumo critico es el volatil y el contexto del mercado geografico o sectorial.
    3. "baseline_revenue": Ingresos mensuales promedio base del negocio en condiciones normales. Debe ser un numero float.
    4. "baseline_cogs": Costo mensual de ventas base (COGS) promedio del negocio. Debe ser un numero float.
    5. "raw_material_share": Proporcion de la materia prima/insumo critico dentro del costo de ventas total (ej. 0.35 si representa el 35% del COGS). Debe ser un float entre 0.05 y 0.90.
    6. "fixed_costs": Gastos fijos mensuales promedio (renta, nomina fija, servicios, seguros). Debe ser un numero float.
    7. "raw_material_risk_prob": Probabilidad de que el precio del insumo critico aumente bruscamente debido a factores de riesgo (clima, logistica, aranceles) (ej. 0.30 para un 30% de probabilidad). Debe ser un float entre 0.0 y 1.0.
    8. "raw_material_risk_increase": Porcentaje estimado de aumento en el costo del insumo en caso de crisis (ej. 0.50 si sube un 50% de precio). Debe ser un float entre 0.0 y 2.0.
    9. "low_season_contraction_mean": Contraccion promedio esperada de la demanda general en los 3 meses de temporada baja (ej. 0.20 para una reduccion del 20% en ventas). Debe ser un float entre 0.0 y 1.0.
    10. "low_season_contraction_std": Volatilidad o desviacion estandar de la contraccion de la demanda en temporada baja. Valor sugerido entre 0.02 y 0.15 (por defecto 0.05).
    11. "high_season_spike_mean": Repunte o incremento promedio esperado de la demanda en el mes de temporada alta (ej. 0.60 para un incremento de 60%). Debe ser un float entre 0.0 y 2.0.
    12. "high_season_spike_std": Volatilidad o desviacion estandar del incremento en temporada alta. Valor sugerido entre 0.05 y 0.25 (por defecto 0.10).
    
    Contenido del documento:
    ---
    {document_text}
    ---
    
    Devuelve UNICAMENTE el objeto JSON plano solicitado. No agregues formatos de código como ```json ... ```, solo la cadena JSON limpia.
    """
    
    # Sanitize the prompt string to prevent ascii encoding issues in API clients
    prompt = clean_to_ascii(prompt)
    
    # Use the passed model name
    model = genai.GenerativeModel(model_name)
    
    # Request JSON response type natively
    response = model.generate_content(
        prompt,
        generation_config={"response_mime_type": "application/json"}
    )
    
    try:
        data = json.loads(response.text)
        return data
    except Exception as e:
        # Fallback to cleaning markdown blocks if any
        cleaned_text = response.text.replace("```json", "").replace("```", "").strip()
        return json.loads(cleaned_text)

